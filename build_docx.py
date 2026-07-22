#!/usr/bin/env python3
"""Render report.md to .docx matching the reference style profile.

Style:
  body:    Arial 11pt
  H1:      Arial 20pt black bold
  H2:      Arial 16pt
  H3:      Arial 14pt #434343
  H4..H6:  Arial 12/11/11pt #666666
  margins: 1 inch
  code blocks: white background, thin black border, 1-cell table, Roboto Mono 9pt
  tables: black header row + white text, white body, thin black borders, cell margin 0.1cm
  images: scaled to fit page width (6.5 in)
"""
import re
from pathlib import Path

import markdown
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from htmldocx import HtmlToDocx

ROOT = Path(__file__).parent
SRC = ROOT / "report.md"
OUT = ROOT / "docs" / "SENG326-Task2-Report-Team13.docx"

BODY_FONT = "Arial"
CODE_FONT = "Consolas"
HEADING_COLOR = "17365D"
H_COLORS = {1: HEADING_COLOR, 2: HEADING_COLOR, 3: HEADING_COLOR,
            4: HEADING_COLOR, 5: HEADING_COLOR, 6: HEADING_COLOR}
H_SIZES = {1: 20, 2: 16, 3: 14, 4: 12, 5: 11, 6: 11}

CODE_TOKEN = "@@@CODEBLOCK_{idx}@@@"
PAGE_WIDTH_IN = 6.5  # 8.5" page - 2x 1" margins
PAGE_HEIGHT_IN = 9.0  # 11" - 2x 1" margins


def set_run_font(run, name=BODY_FONT, size_pt=None, color_hex=None, bold=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color_hex is not None:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    if bold is not None:
        run.font.bold = bold


def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=100, left=100, bottom=100, right=100):
    """Margins in twentieths of a point (twips). 100 twips ≈ 0.07cm."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for direction, val in [("top", top), ("left", left),
                            ("bottom", bottom), ("right", right)]:
        node = OxmlElement(f"w:{direction}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_borders(cell, color="000000", size="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def configure_styles(doc):
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), BODY_FONT)

    for level, size in H_SIZES.items():
        s = styles[f"Heading {level}"]
        s.font.name = BODY_FONT
        s.font.size = Pt(size)
        s.font.bold = True
        if H_COLORS[level]:
            s.font.color.rgb = RGBColor.from_string(H_COLORS[level])


def add_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SENG 326 — Software Architecture")
    set_run_font(r, size_pt=26, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Task 2 — Architectural Refactoring of Checkstyle 13.2.0")
    set_run_font(r, size_pt=15, color_hex="666666")

    doc.add_paragraph()
    doc.add_paragraph()

    target = doc.add_paragraph()
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = target.add_run("Target architecture: Pipe-and-Filter")
    set_run_font(r, size_pt=14, bold=True)

    doc.add_paragraph()

    for line in [
        "KARIM HARIRI – 22050941008",
        "MOHAMED ATTIA EID ATTIA EID – 22050941017",
        "ZAID HARDAN – 22050941005",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run_font(r, size_pt=12)

    doc.add_page_break()


def strip_cover_from_md(md_text):
    lines = md_text.splitlines()
    for i, line in enumerate(lines):
        if line.strip().lower().startswith("## executive summary"):
            return "\n".join(lines[i:])
    return md_text


def extract_code_blocks(md_text):
    """Replace ```...``` fenced code with markers; return (text, codes_list)."""
    codes = []

    def repl(m):
        idx = len(codes)
        codes.append(m.group(2))
        return f"\n\n{CODE_TOKEN.format(idx=idx)}\n\n"

    pat = re.compile(r"```([A-Za-z0-9_-]*)\n(.*?)```", re.DOTALL)
    new = pat.sub(repl, md_text)
    return new, codes


def insert_code_box(doc, paragraph, code_text):
    """Replace a marker paragraph with a 1-cell bordered table containing code.
    Empty box with thin black border, white background, monospace contents."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    # apply borders + padding + white fill
    set_cell_borders(cell, color="000000", size="6")
    set_cell_shading(cell, "FFFFFF")
    set_cell_margins(cell, top=180, left=240, bottom=180, right=240)
    # set table column width to ~ page width
    table.columns[0].width = Inches(PAGE_WIDTH_IN)
    cell.width = Inches(PAGE_WIDTH_IN)
    # remove default paragraph
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    for i, line in enumerate(code_text.rstrip("\n").splitlines()):
        if i > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line if line else " ")
        set_run_font(r, name=CODE_FONT, size_pt=9)
    # move our newly added table into position by swapping XML nodes
    tbl_el = table._element
    target_el = paragraph._element
    target_el.addnext(tbl_el)
    target_el.getparent().remove(target_el)


def style_data_tables(doc):
    """Apply black header / white body / 0.1cm padding to all tables produced
    by htmldocx (skip code-box tables we already styled)."""
    for table in doc.tables:
        # detect 1x1 code box: already styled white fill -> skip
        if len(table.rows) == 1 and len(table.columns) == 1:
            continue
        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                set_cell_borders(cell, color="000000", size="6")
                set_cell_margins(cell, top=56, left=56, bottom=56, right=56)
                if r_idx == 0:
                    set_cell_shading(cell, "000000")
                else:
                    set_cell_shading(cell, "FFFFFF")
                for p in cell.paragraphs:
                    for run in p.runs:
                        if r_idx == 0:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.bold = True
                        else:
                            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                        if not run.font.name:
                            set_run_font(run, name=BODY_FONT, size_pt=10)
                        else:
                            run.font.name = BODY_FONT
                            run.font.size = Pt(10)


def fit_images_to_page(doc):
    """Scale every inline image so it fits within the printable page width AND
    is no taller than the printable page height."""
    EMU_PER_INCH = 914400
    max_w = int(PAGE_WIDTH_IN * EMU_PER_INCH)
    max_h = int(PAGE_HEIGHT_IN * EMU_PER_INCH)
    for shape in doc.inline_shapes:
        w, h = shape.width, shape.height
        if w == 0 or h == 0:
            continue
        scale_w = max_w / w if w > max_w else 1
        scale_h = max_h / h if h > max_h else 1
        scale = min(scale_w, scale_h, 1.0)
        if scale < 1.0:
            shape.width = int(w * scale)
            shape.height = int(h * scale)


# ============ build ============

raw_md = strip_cover_from_md(SRC.read_text(encoding="utf-8"))
md_no_code, code_blocks = extract_code_blocks(raw_md)

html = markdown.markdown(
    md_no_code,
    extensions=["fenced_code", "tables", "toc", "sane_lists"],
)

doc = Document()
configure_styles(doc)
add_cover(doc)

converter = HtmlToDocx()
converter.table_style = "Table Grid"  # plain bordered; we override cells next
converter.add_html_to_document(html, doc)

# Strip cnfStyle from every row/cell so banded-row formatting can't override us
for tbl in doc.tables:
    for row in tbl.rows:
        for el in row._tr.iter(qn("w:cnfStyle")):
            el.getparent().remove(el)
        for cell in row.cells:
            for el in cell._tc.iter(qn("w:cnfStyle")):
                el.getparent().remove(el)

# Replace each marker paragraph with a code-box table
marker_re = re.compile(r"^@@@CODEBLOCK_(\d+)@@@$")
to_process = []
for p in doc.paragraphs:
    m = marker_re.match(p.text.strip())
    if m:
        to_process.append((p, int(m.group(1))))
for p, idx in to_process:
    insert_code_box(doc, p, code_blocks[idx])

# style data tables (black header / white body)
style_data_tables(doc)

# Re-apply Arial body font + heading run color/bold + italic Figure captions
for p in doc.paragraphs:
    style_name = (p.style.name or "") if p.style else ""
    if style_name.startswith("Heading"):
        for r in p.runs:
            set_run_font(r, name=BODY_FONT,
                         color_hex=HEADING_COLOR, bold=True)
    elif p.text.strip().startswith("Figure:") and len(p.runs) >= 1:
        # caption paragraph — italic, smaller
        for r in p.runs:
            set_run_font(r, name=BODY_FONT, size_pt=9)
            r.italic = True
            r.bold = False
            r.font.color.rgb = RGBColor.from_string("595959")
    else:
        for r in p.runs:
            if not r.font.name:
                set_run_font(r, name=BODY_FONT)

# Fit all images to page
fit_images_to_page(doc)

doc.save(str(OUT))
print(f"wrote {OUT}")
